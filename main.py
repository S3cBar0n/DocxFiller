from docx.shared import Pt
from datetime import date
from docx import Document
import pyad.pyadutils
import pyad.adquery
import pyad.aduser
import datetime
import os.path
import logging
import utils
import pytz


# Function to determine which set of password reset instructions to give the user
def find_job_title(title):
    now = datetime.datetime.now()
    current_time = now.strftime("%H:%M:%S")
    thinclient = ['Receptionist', 'Janitor']
    laptop = ['Manager', 'Vice President']
    desktop = ['Technical Support', 'Systems Administrator']

    all_titles = thinclient + laptop + desktop

    for item in all_titles:
        if item in thinclient and (title in item or item in title):
            deviceType = 'Once you have logged into your device, you will want to press & ' \
                         'hold these keys in this order: Ctrl + Alt + Del' \
                         '\n\n' \
                         'Note: You will not use/press the ' \
                         '“+” key for the instructions above, it serves merely to indicate you will want to ' \
                         'press the following keys all together.'
            return deviceType

        elif item in laptop and (title in item or item in title):
            deviceType = 'Once you have logged into your device, you will want to press & ' \
                         'hold these keys in this order: Ctrl + Alt + Del' \
                         '\n\n' \
                         'Note: You will not use/press the ' \
                         '“+” key for the instructions above, it serves merely to indicate you will want to ' \
                         'press the following keys all together.'
            return deviceType

        elif item in desktop and (title in item or item in title):
            deviceType = 'Once you have logged into your device, you will want to press & ' \
                         'hold these keys in this order: Ctrl + Alt + Del' \
                         '\n\n' \
                         'Note: You will not use/press the ' \
                         '“+” key for the instructions above, it serves merely to indicate you will want to ' \
                         'press the following keys all together.'
            return deviceType
        else:
            pass
    else:
        # Sets the password reset instructions to the below string if the title has not been seen before.
        deviceType = "-WARNING- UNABLE TO DETERMINE PASSWORD INSTRUCTIONS BASED OFF OF TITLE, PLEASE POPULATE THIS " \
                     "FIELD!"
        logging.warning(f"{current_time} - FAILED TO DETERMINE PASSWORD CHANGE INSTRUCTIONS FOR USER.")
        return deviceType


# Our main function to pull Active Directory information and create letters.
def main():
    # Variables used to help calculate and target users who have been created in the last 30 days
    today = datetime.datetime.utcnow().replace(tzinfo=pytz.UTC)
    days = datetime.timedelta(days=30)
    alpha = today - days

    # Variable to store which users we created letters for - Used later in the email.
    users = []

    # Time/Date variables to get the current date for the logfile name.
    now = datetime.datetime.now()
    current_time = now.strftime("%H:%M:%S")
    logging.basicConfig(filename=f'C:\\LOCATION_TO_STORE_LOG\\{date.today()}.log',
                        level=logging.DEBUG)
    try:
        logging.info(f"{current_time} - Gathering Active Directory Information...")
        # Launching the AD query service
        q = pyad.adquery.ADQuery()

        # Giving the query service our search and return criteria
        q.execute_query(
            attributes=["whenCreated", "cn", "employeeID", "pwdLastSet", "description", "displayName", "SamAccountName",
                        "mail",
                        "telephoneNumber",
                        "facsimileTelephoneNumber",
                        "Useraccountcontrol"],
            where_clause="objectClass = 'user'",
            base_dn="OU=Users, DC=DOMAIN_NAME, DC=local"
        )
    except Exception as e:
        logging.error(f"{current_time} - {e}!")
        logging.error(f"{current_time} - Failed to initialize the Active Directory Query Service! Quitting!")
        exit()

    logging.info(f"{current_time} - Locating users created within the last two weeks...")
    # Checks to see if any users were pulled from the search, then filters out devices, terminations, service accounts.
    # Also checks to see if the user pulled has been created in the last 30 days
    for row in q.get_results():
        if "$" not in str(row['displayName']):
            if row['whenCreated'] >= alpha:  # This is checking to make sure the New Hire at maximum 30 days of age
                if row['displayName'] is not None and row['telephoneNumber'] is not None:
                    # This determines whether the user's account is disabled or terminated.
                    if (not row["Useraccountcontrol"] / 2 % 2 != 0) or ("TERMINATED" not in str(row['displayName'])):
                        if row['employeeID'] != 1234567890:
                            # Checks to see if the user has a fax number, chooses the template accordingly
                            if row['facsimileTelephoneNumber'] is not None:
                                doc = Document("Template_with_fax.docx")
                                doc2 = Document("Template_with_fax_manager.docx")
                                logging.info(
                                    f"{current_time} - Preparing {row['displayName']}'s letter with Fax Fields...")
                            else:
                                doc = Document("Template.docx")
                                doc2 = Document("Template_manager.docx")
                                logging.info(
                                    f"{current_time} - Preparing {row['displayName']}'s letter without Fax Fields...")

                            file_path = "YOUR_FILE_PATH_TO_PLACE_THE_FINISHED_LETTER"
                            delivered_path = "FILE_PATH_IF_TO_SEE_IF_ANY_LETTERS_HAVE_BEEN_CREATED_IN_THE_PAST"
                            # Checks to see if the user already has a letter that has been delivered
                            if os.path.isdir(f"{delivered_path}\\{row['displayName']}"):
                                logging.warning(
                                    f"{current_time} - {row['displayName']} already has a folder in the Delivered Letters folder... Skipping user!")
                                print(f"{row['displayName']}'s folder exists in delivered letters... Skipping...")
                                continue
                            else:
                                pass

                            # Checks to see if the letters exist, and if not creates each one even if one is missing
                            if os.path.isdir(f"{file_path}\\{row['displayName']}") and os.path.isfile(
                                    f"{file_path}\\{row['displayName']}\\" + f"New Hire Letter {row['displayName']}.docx") and os.path.isfile(
                                    f"{file_path}\\{row['displayName']}\\" + f"New Hire Letter {row['displayName']} Manager.docx"):
                                logging.warning(
                                    f"{current_time} - {row['displayName']} already has a folder, and 2 letters... Skipping user!")
                                print(f"New Hire Letter {row['displayName']}.docx exists... Skipping...")
                                continue
                            elif os.path.isdir(f"{file_path}\\{row['displayName']}") is False:
                                logging.info(f"{current_time} - Creating {row['displayName']} folder...")
                                os.mkdir(f"{file_path}\\{row['displayName']}")
                                pass
                            else:
                                logging.warning(
                                    f"{current_time} - {row['displayName']} already has a directory and 1 letter, attempted to "
                                    f"generate other letter. Please review this individual!")
                                pass

                            # Takes the pwdLastSet and whenCreated attribute and converts them to a date that is
                            # readable for the program to do its calculations
                            last_set = pyad.pyadutils.convert_bigint(row['pwdLastSet'])
                            created = row['whenCreated'].strftime('%Y-%m-%d %H:%M:%S')
                            created = datetime.datetime.strptime(created, "%Y-%m-%d %H:%M:%S")
                            timestamp = last_set
                            last_set = datetime.datetime(1601, 1, 1) + datetime.timedelta(seconds=timestamp / 10000000)
                            last_set = last_set.strftime('%Y-%m-%d %H:%M:%S')
                            last_set = datetime.datetime.strptime(last_set, "%Y-%m-%d %H:%M:%S")
                            elapsed_time = last_set - created

                            # This is checking to see if the user's Creation Time and Password Last Set are within 5
                            # seconds of each other and changes the password if they are.
                            if elapsed_time < datetime.timedelta(seconds=5):
                                password = utils.reset_user_password(row['cn'])
                                logging.info(f"{current_time} - {row['displayName']}'s password has been changed.")
                            else:
                                print(f"{row['displayName']} does not need a password change\n")
                                logging.warning(
                                    f"{current_time} - {row['displayName']}'s password has NOT been changed.")
                                password = "WARNING - COULD NOT SET PASSWORD"
                                pass

                            # Initializing the styles we need for the docx
                            style = doc.styles['Normal']
                            font = style.font
                            font.name = 'Calibri'
                            font.size = Pt(11)

                            style2 = doc2.styles['Normal']
                            font2 = style2.font
                            font2.name = 'Calibri'
                            font2.size = Pt(11)

                            # From here to line 235 is assigning the values we need to certain fields so when its
                            # saved they will be in their proper locations.
                            description = str(row["description"])
                            for sym in (("'", ""), (",", ""), ("(", ""), (")", "")):
                                description = description.replace(*sym)
                            deviceType = find_job_title(description.replace(",", ""))

                            user_table = doc.tables[0]
                            user_table.cell(0, 1).text = row['SamAccountName']
                            user_table.cell(1, 1).text = password

                            man_table = doc2.tables[0]
                            man_table.cell(0, 1).text = row['SamAccountName']
                            man_table.cell(1, 1).text = password

                            user_table = doc.tables[1]
                            user_table.cell(0, 0).text = deviceType

                            man_table = doc2.tables[1]
                            man_table.cell(0, 0).text = deviceType

                            user_table = doc.tables[2]
                            user_table.cell(0, 1).text = row['mail']

                            man_table = doc2.tables[2]
                            man_table.cell(0, 1).text = row['mail']

                            user_table = doc.tables[3]
                            user_table.cell(0, 1).text = row['telephoneNumber']
                            if row['facsimileTelephoneNumber'] is not None:
                                user_table.cell(2, 1).text = row['facsimileTelephoneNumber']
                            else:
                                pass

                            man_table = doc2.tables[3]
                            man_table.cell(0, 1).text = row['telephoneNumber']
                            if row['facsimileTelephoneNumber'] is not None:
                                man_table.cell(2, 1).text = row['facsimileTelephoneNumber']
                            else:
                                pass

                            user_table = doc.tables[4]
                            user_table.cell(0, 1).text = row['mail']

                            man_table = doc2.tables[4]
                            man_table.cell(0, 1).text = row['mail']

                            user_table = doc.tables[5]
                            user_table.cell(1, 1).text = row['SamAccountName']

                            user_table = doc.tables[6]
                            user_table.cell(1, 1).text = row['SamAccountName']

                            user_table = doc.tables[7]
                            user_table.cell(0, 1).text = row['SamAccountName']

                            # Attempting to save the two documents based off of the values we assigned above.
                            try:
                                logging.info(f"{current_time} - Beginning User and Manager letter generation...")
                                doc.save(
                                    f"{file_path}\\{row['displayName']}\\" + f"New Hire Letter {row['displayName']}.docx")
                                doc2.save(
                                    f"{file_path}\\{row['displayName']}\\" + f"New Hire Letter {row['displayName']} Manager.docx")
                                logging.info(
                                    f"{current_time} - Successfully created User and Manager letter for {row['displayName']}.")
                                users.append(row['displayName'])
                            except Exception as e:
                                logging.error(f"{current_time} - {e}!")
                                logging.error(f"{current_time} - Failed to create Letters for {row['displayName']}!")
                        else:
                            logging.warning(f"Skipping {row['displayName']}, detected service account.")
                            print(f"Skipping {row['displayName']}, detected service account.")
                            continue
                    else:
                        logging.warning(f"{current_time} - Skipping {row['displayName']}, they are disabled.")
                        print(f"Skipping {row['displayName']}, they are disabled.\n")
                        continue
                else:
                    logging.info(f"Skipping {row['displayName']}, their AD Profile is not yet filled out..")
                    print(f"{row['displayName']}'s AD Profile is not yet filled out..")
                    continue
            else:
                continue
        else:
            continue

    # Checks to see if any users had a letter generated and sends out the Log and results email.

    logging.info(f"{current_time} - End of current run.\n")
    logging.shutdown()

    if users:
        users.sort()
        utils.send_log(users)
    else:
        pass


if __name__ == '__main__':
    main()
